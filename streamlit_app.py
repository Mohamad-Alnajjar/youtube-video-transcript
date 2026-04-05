import io
import logging
import re
import anthropic
from urllib.parse import urlparse, parse_qs

import streamlit as st
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
from youtube_transcript_api import (
    CouldNotRetrieveTranscript,
    NoTranscriptFound,
    TranscriptsDisabled,
    YouTubeTranscriptApi,
)
from youtube_transcript_api.proxies import WebshareProxyConfig

# ── Logging ───────────────────────────────────────────────────────────────────
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Register built-in Japanese CID font once at module load
pdfmetrics.registerFont(UnicodeCIDFont("HeiseiMin-W3"))

# ── Secrets ───────────────────────────────────────────────────────────────────

def _secret(key: str) -> str:
    try:
        return st.secrets.get(key, "") or ""
    except Exception:
        return ""

# ── YouTube helpers ───────────────────────────────────────────────────────────

def extract_video_id(url: str) -> str:
    url = url.strip()
    if "youtube.com" in url or "youtu.be" in url:
        parsed = urlparse(url)
        qs = parse_qs(parsed.query)
        if "v" in qs:
            return qs["v"][0]
        path_id = parsed.path.lstrip("/").split("/")[0]
        if path_id:
            return path_id
        raise ValueError("Could not extract video ID from URL.")
    if re.match(r"^[A-Za-z0-9_-]{11}$", url):
        return url
    raise ValueError("Invalid YouTube URL or video ID.")


def _yt_api() -> YouTubeTranscriptApi:
    proxy_user = _secret("WEBSHARE_PROXY_USERNAME")
    proxy_pass = _secret("WEBSHARE_PROXY_PASSWORD")
    if proxy_user and proxy_pass:
        return YouTubeTranscriptApi(
            proxy_config=WebshareProxyConfig(
                proxy_username=proxy_user,
                proxy_password=proxy_pass,
            )
        )
    return YouTubeTranscriptApi()


def list_available_languages(video_id: str) -> dict[str, str]:
    try:
        transcript_list = _yt_api().list(video_id)
        langs = {}
        for t in transcript_list:
            label = t.language + (" (auto-generated)" if t.is_generated else "")
            langs[t.language_code] = label
        return langs
    except TranscriptsDisabled:
        raise RuntimeError("Subtitles are disabled for this video.")
    except Exception as exc:
        err_str = str(exc)
        logger.exception("Failed to list transcripts for %s", video_id)
        if "RequestBlocked" in type(exc).__name__ or "blocked" in err_str.lower():
            raise RuntimeError(
                "YouTube is blocking requests from this server's IP. "
                "Add Webshare proxy credentials to Streamlit secrets — see the sidebar."
            ) from exc
        raise RuntimeError(f"Could not retrieve subtitle information: {exc}") from exc


def fetch_transcript(video_id: str, lang_code: str) -> list:
    try:
        transcript_list = _yt_api().list(video_id)
        try:
            transcript = transcript_list.find_manually_created_transcript([lang_code])
        except NoTranscriptFound:
            transcript = transcript_list.find_generated_transcript([lang_code])
        return transcript.fetch()
    except NoTranscriptFound:
        raise RuntimeError(f"No transcript found for '{lang_code}'. Try a different language.")
    except TranscriptsDisabled:
        raise RuntimeError("Subtitles are disabled for this video.")
    except CouldNotRetrieveTranscript as exc:
        raise RuntimeError(f"Could not retrieve transcript: {exc}") from exc
    except Exception as exc:
        logger.exception("Unexpected error fetching transcript")
        raise RuntimeError(f"Failed to fetch transcript: {exc}") from exc


def format_seconds(seconds: float) -> str:
    total = int(seconds)
    h, remainder = divmod(total, 3600)
    m, s = divmod(remainder, 60)
    return f"{h}:{m:02d}:{s:02d}" if h else f"{m}:{s:02d}"


def entries_to_raw_text(entries: list) -> str:
    return "".join(
        (getattr(e, "text", None) or e.get("text", "")).strip()
        for e in entries
    )

# ── Claude formatting ─────────────────────────────────────────────────────────

STUDY_SYSTEM_PROMPT = """You are a Japanese language study assistant.
You will receive raw auto-generated Japanese transcript text that has no punctuation and is split at arbitrary time boundaries, making it hard to read.

Your job is to reformat it into a clean, study-friendly document.

Rules:
- Split the text into proper sentences based on meaning and natural speech boundaries.
- Add appropriate Japanese punctuation: 。？！、 where they belong.
- Each sentence goes on its own line.
- Do NOT add furigana, romaji, or any other annotations unless asked.
- Do NOT skip or summarise any content — include everything.
- Do NOT add any headings, labels, or commentary of your own.
- Output only the formatted transcript, nothing else.

Example input:
ネット上で度々起こるある議論それ
が書かないで覚える方が効率
的アプリで勉強した方が効率いい
じゃん何回も書いた方が手が
覚える書いて覚えないとすぐ忘れ
ちゃう
Example output:
ネット上で度々起こるある議論,それが書かないで覚える方が効率的

アプリで勉強した方が効率いいじゃん

何回も書いた方が手が覚える

書いて覚えないとすぐ忘れちゃう
"""


def format_with_claude(raw_text: str) -> str:
    api_key = _secret("ANTHROPIC_API_KEY")
    if not api_key:
        raise RuntimeError("No Anthropic API key found. Add ANTHROPIC_API_KEY to your Streamlit secrets.")

    client = anthropic.Anthropic(api_key=api_key)
    placeholder = st.empty()
    full_response = ""

    with client.messages.stream(
        model="claude-haiku-4-5",
        max_tokens=8000,
        system=STUDY_SYSTEM_PROMPT,
        messages=[{"role": "user", "content": raw_text}],
    ) as stream:
        for text in stream.text_stream:
            full_response += text
            placeholder.markdown(f"```\n{full_response}▌\n```")

    placeholder.empty()
    return full_response

# ── DOCX builders ─────────────────────────────────────────────────────────────

def build_raw_docx(entries: list, title: str, include_timestamps: bool) -> bytes:
    doc = Document()
    heading = doc.add_heading(title, level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.add_paragraph()

    for entry in entries:
        text = (getattr(entry, "text", None) or entry.get("text", "")).strip()
        start = getattr(entry, "start", None) or entry.get("start", 0)
        if include_timestamps:
            p = doc.add_paragraph()
            ts = p.add_run(f"[{format_seconds(start)}]  ")
            ts.bold = True
            ts.font.size = Pt(9)
            p.add_run(text)
        else:
            doc.add_paragraph(text)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_study_docx(formatted_text: str, title: str) -> bytes:
    doc = Document()
    heading = doc.add_heading(title + " — Study Edition", level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.add_paragraph()

    for line in formatted_text.strip().splitlines():
        line = line.strip()
        if not line:
            doc.add_paragraph()
            continue
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.size = Pt(13)
        p.paragraph_format.space_after = Pt(6)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()

# ── PDF builder ───────────────────────────────────────────────────────────────

def build_study_pdf(formatted_text: str, title: str) -> bytes:
    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=A4,
        leftMargin=20 * mm,
        rightMargin=20 * mm,
        topMargin=20 * mm,
        bottomMargin=20 * mm,
    )

    title_style = ParagraphStyle(
        "JpTitle",
        fontName="HeiseiMin-W3",
        fontSize=16,
        leading=24,
        spaceAfter=12,
    )
    body_style = ParagraphStyle(
        "JpBody",
        fontName="HeiseiMin-W3",
        fontSize=13,
        leading=22,
        spaceAfter=8,
    )

    story = [Paragraph(title + " — Study Edition", title_style), Spacer(1, 6 * mm)]

    for line in formatted_text.strip().splitlines():
        line = line.strip()
        if not line:
            story.append(Spacer(1, 4 * mm))
        else:
            story.append(Paragraph(line, body_style))

    doc.build(story)
    return buf.getvalue()

# ── UI ────────────────────────────────────────────────────────────────────────

def main():
    st.set_page_config(page_title="YouTube → Japanese Study Doc", page_icon="📄")
    st.title("📄 YouTube → Japanese Study Doc")
    st.write("Fetch a YouTube transcript and reformat it into clean, study-ready Japanese sentences.")

    # ── Sidebar ────────────────────────────────────────────────────────────
    with st.sidebar:
        st.header("⚙️ Setup")

        if _secret("WEBSHARE_PROXY_USERNAME") and _secret("WEBSHARE_PROXY_PASSWORD"):
            st.success("✅ YouTube proxy configured")
        else:
            st.warning("⚠️ No YouTube proxy")
            st.markdown("""
Running locally? No proxy needed.
On Streamlit Cloud, add to **Settings → Secrets**:
```toml
WEBSHARE_PROXY_USERNAME = "..."
WEBSHARE_PROXY_PASSWORD = "..."
```
""")

        if _secret("ANTHROPIC_API_KEY"):
            st.success("✅ Anthropic API key configured")
        else:
            st.info("""ℹ️ Add your Anthropic API key to Secrets:
```toml
ANTHROPIC_API_KEY = "sk-ant-..."
```
Get a key at [console.anthropic.com](https://console.anthropic.com)
""")

    # ── Step 1: URL ────────────────────────────────────────────────────────
    video_url = st.text_input("🎥 YouTube Video URL or Video ID")

    video_id = None
    available_langs: dict[str, str] = {}

    if video_url.strip():
        try:
            video_id = extract_video_id(video_url)
        except ValueError as exc:
            st.error(f"❌ {exc}")

    if video_id:
        with st.spinner("Fetching available subtitle languages…"):
            try:
                available_langs = list_available_languages(video_id)
            except RuntimeError as exc:
                st.error(f"❌ {exc}")
                video_id = None

    # ── Step 2: Language ───────────────────────────────────────────────────
    lang_code = None
    if available_langs:
        lang_options = {v: k for k, v in available_langs.items()}
        chosen_label = st.selectbox("🌐 Subtitle language", options=list(lang_options.keys()))
        lang_code = lang_options[chosen_label]

    # ── Step 3: Options ────────────────────────────────────────────────────
    col1, col2 = st.columns(2)
    with col1:
        include_timestamps = st.checkbox("⏱ Timestamps in raw doc", value=True)
    with col2:
        study_mode = st.checkbox("🎓 Generate study version", value=True)

    # ── Step 4: Generate ───────────────────────────────────────────────────
    if st.button("Generate Transcript", disabled=(not lang_code)):
        with st.spinner("Fetching transcript…"):
            try:
                entries = fetch_transcript(video_id, lang_code)
            except RuntimeError as exc:
                st.error(f"❌ {exc}")
                return

        st.success(f"✅ Fetched {len(entries)} transcript entries.")
        safe_title = re.sub(r'[\\/*?:"<>|]', "_", f"Video_{video_id}")

        # ── Raw transcript ─────────────────────────────────────────────
        with st.expander("📝 Raw Transcript", expanded=not study_mode):
            preview_lines = []
            for e in entries:
                start = getattr(e, "start", None) or e.get("start", 0)
                text = (getattr(e, "text", None) or e.get("text", "")).strip()
                ts = f"[{format_seconds(start)}] " if include_timestamps else ""
                preview_lines.append(f"{ts}{text}")
            st.text_area("Raw transcript text", "\n".join(preview_lines), height=250, label_visibility="collapsed")

        raw_docx = build_raw_docx(entries, safe_title, include_timestamps)
        st.download_button(
            label="📥 Download raw .docx",
            data=raw_docx,
            file_name=f"{safe_title}_raw.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

        # ── Study version ──────────────────────────────────────────────
        if study_mode:
            st.divider()
            st.subheader("🎓 Study Version")
            st.caption("Claude is reformatting the transcript into clean, punctuated Japanese sentences…")

            raw_text = entries_to_raw_text(entries)

            try:
                formatted = format_with_claude(raw_text)
            except Exception as exc:
                st.error(f"❌ Claude formatting failed: {exc}")
                return

            st.text_area("📖 Formatted Transcript", formatted, height=400)

            # Download buttons side by side
            dl_col1, dl_col2 = st.columns(2)

            with dl_col1:
                study_docx = build_study_docx(formatted, safe_title)
                st.download_button(
                    label="📥 Download study .docx",
                    data=study_docx,
                    file_name=f"{safe_title}_study.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )

            with dl_col2:
                study_pdf = build_study_pdf(formatted, safe_title)
                st.download_button(
                    label="📥 Download study .pdf",
                    data=study_pdf,
                    file_name=f"{safe_title}_study.pdf",
                    mime="application/pdf",
                )


if __name__ == "__main__":
    main()